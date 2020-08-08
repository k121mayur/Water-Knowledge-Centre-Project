from tkinter import *
from tkinter import ttk
from tkinter import messagebox
from PIL import ImageTk, Image
import sqlite3
import datetime
from docx import *
import os
from tkcalendar import Calendar

mk = os.access("word/bill", mode=0o777)

if mk is False:
    os.makedirs("word/bill")

date = datetime.date.today()
formatted_date = date.strftime('%d %B %Y')

window = Tk()
window.title('Water Quality Testing Laboratory')
window.iconbitmap('')

# Database connection
connection = sqlite3.connect('WKC.db')
cur = connection.cursor()

cur.execute('''CREATE TABLE IF NOT EXISTS main_data ("Id" INTEGER NOT NULL PRIMARY KEY AUTOINCREMENT UNIQUE, 
"customer" TEXT NOT NULL, "sample_number" INTEGER NOT NULL UNIQUE, "sample_drawn_by" TEXT, "sample_drawn_date" TEXT, 
"sample_reached_lab" TEXT, "test_start_date" TEXT, "test_end_date" TEXT, "sample_reference" TEXT, "village" TEXT, 
"source_type" TEXT, "location_of_source" TEXT, "appearance" TEXT, "colour" TEXT, "odour" TEXT, "turbidity" TEXT, 
"electrical_conductivity" TEXT, "total_dissolved_solids" TEXT, "total_solids" TEXT, "total_suspended_solids" TEXT, 
"ph" TEXT, "ph_alkalinity" TEXT, "total_alkalinity" TEXT, "total_hardness" TEXT, "ca" TEXT, "mg" TEXT, "fe" TEXT, 
"na" TEXT, "k" TEXT, "nh3" TEXT, "no2" TEXT, "no3" TEXT, "cl" TEXT, "f" TEXT, "so4" TEXT, "po4" TEXT, "tids_test" TEXT, 
"do" TEXT, "bod" TEXT, "chemical_oxygen_demand" TEXT, "fecal_coliform" TEXT, "comments" TEXT, "date" TEXT)''')


# First / Start screen
def start():
    # for back button
    try:
        next_button1.grid_forget()
        back_to_start.grid_forget()
    except:
        pass

    # Checking last entry is Available or not.
    try:
        cur.execute('SELECT max(id) FROM main_data')
        max_id = cur.fetchall()[0][0]
    except:
        max_id = None

    if max_id is None:
        pass
    else:
        cur.execute('''SELECT * FROM main_data WHERE id = ?''', (max_id,))
        all = cur.fetchall()
        customer_name = all[0][1]
        sample_no = all[0][2]
        sample_drawn_by = all[0][3]
        drawn_date = all[0][4]
        reached_date = all[0][5]
        start_date = all[0][6]
        end_date = all[0][7]
        sample_ref = all[0][8]
        village = all[0][9]
        source = all[0][10]
        location = all[0][11]

    # Exit a confirmation window if Open else pass
    try:
        top.destroy()
    except:
        pass

    # Erase a Physical Parameters Window
    try:
        frame2.grid_forget()
        next_button1.grid_forget()
    except:
        pass

    # Erase a Chemical Parameters Window
    try:
        chemical_half.grid_forget()
        next_comments.grid_forget()
    except:
        pass

    leb1 = Label(window, text="Water Quality Testing Laboratory", fg='black')
    leb1.grid(row=0, column=0, sticky=W, padx=20)  # columnspan = 0)

    leb2 = Label(window, text="Water Knowledge Centre (WKC)", fg='black')
    leb2.grid(row=1, column=0, sticky=W, padx=20)  # , columnspan = 0)

    leb3 = Label(window, text="The DHAN Academy (TDA), Madurai", fg='black')
    leb3.grid(row=3, column=0, sticky=W, padx=20)  # ,columnspan = 0 )

    leb4 = Label(window, text="TEST ANALYSIS REPORT", fg='black', bg='grey', font=32)
    leb4.grid(row=4, column=0, ipady=10, pady=10, columnspan=5, sticky=EW, )

    data = Image.open('tda_logo_150px.jpg')
    render = ImageTk.PhotoImage(data)
    leb5 = Label(image=render, text='mayur', bg='grey', fg='red')
    leb5.image = render
    leb5.grid(row=0, column=4, rowspan=4, columnspan=2)

    # Sample details
    global frame
    frame = LabelFrame(window, padx=20, pady=20)
    frame.grid(row=5, padx=20, pady=20, columnspan=5)

    leb6 = Label(frame, text='Customer Details:')
    leb6.grid(column=0, row=6, sticky=E, padx=20)

    global entry_customer
    entry_customer = ttk.Entry(frame)
    entry_customer.grid(column=1, row=6, columnspan=4, ipadx=250)

    leb7 = Label(frame, text='Sample Number:')
    leb7.grid(column=0, row=7, sticky=E, padx=20)

    global entry_sampleNumber
    entry_sampleNumber = Entry(frame)
    entry_sampleNumber.grid(column=1, row=7)

    leb8 = Label(frame, text='Sample Drawn By:')
    leb8.grid(column=3, row=7, sticky=E)

    global entry_sampleDrawn
    entry_sampleDrawn = Entry(frame)
    entry_sampleDrawn.grid(column=4, row=7)

    leb9 = Label(frame, text='Sample Drawn Date:')
    leb9.grid(column=0, row=8, sticky=E, padx=5)

    global entry_sampleDrawnDate
    entry_sampleDrawnDate = ttk.Button(frame, text='--Select a date--', command=lambda: datePicker('drawnDate'),
                                       width=18)
    entry_sampleDrawnDate.grid(column=1, row=8, pady=5)

    leb10 = Label(frame, text='Sample Reached Lab on:')
    leb10.grid(column=3, row=8, sticky=E)

    global entry_sampleReached
    entry_sampleReached = ttk.Button(frame, text='--Select a date--', command=lambda: datePicker('reachedDate'),
                                     width=18)
    entry_sampleReached.grid(column=4, row=8, pady=5)

    leb11 = Label(frame, text='Test Start Date:')
    leb11.grid(column=0, row=9, sticky=E, padx=20)

    global entry_testStart
    entry_testStart = ttk.Button(frame, text='--Select a date--', command=lambda: datePicker('testStart'), width=18)
    entry_testStart.grid(column=1, row=9, pady=5)

    leb12 = Label(frame, text='Test End Date:')
    leb12.grid(column=3, row=9, sticky=E)

    global entry_testEnd
    entry_testEnd = ttk.Button(frame, text='--Select a date--', command=lambda: datePicker('testEnd'), width=18)
    entry_testEnd.grid(column=4, row=9, pady=5)

    leb13 = Label(frame, text='Village Name:')
    leb13.grid(column=3, row=10, padx=20, sticky=E)

    global entry_village
    entry_village = Entry(frame)
    entry_village.grid(column=4, row=10)

    leb16 = Label(frame, text='Sample Reference:')
    leb16.grid(column=0, row=10, padx=20, sticky=E)

    global entry_sampleReference
    entry_sampleReference = Entry(frame)
    entry_sampleReference.grid(column=1, row=10)

    leb14 = Label(frame, text="Source Type:")
    leb14.grid(column=0, row=11, sticky=E)

    global entry_sourceType
    entry_sourceType = Entry(frame)
    entry_sourceType.grid(column=1, row=11)

    leb15 = Label(frame, text="Location of Source:")
    leb15.grid(column=3, row=11, sticky=E)

    global entry_location
    entry_location = Entry(frame, text="Village Name:")
    entry_location.grid(column=4, row=11)

    if max_id != None:
        entry_customer.delete(0, END)
        entry_sampleNumber.delete(0, END)
        entry_sampleDrawn.delete(0, END)
        entry_village.delete(0, END)
        entry_sourceType.delete(0, END)
        entry_location.delete(0, END)

        entry_customer.insert(0, customer_name)
        entry_sampleNumber.insert(0, sample_no)
        entry_sampleDrawn.insert(0, sample_drawn_by)
        entry_sampleDrawnDate.configure(text=drawn_date)
        entry_sampleReached.configure(text=reached_date)
        entry_testStart.configure(text=start_date)
        entry_testEnd.configure(text=end_date)
        entry_sampleReference.insert(0, sample_ref)
        entry_village.insert(0, village)
        entry_sourceType.insert(0, source)
        entry_location.insert(0, location)

    global next_button
    next_button = ttk.Button(window, text="Next", command=physical)
    next_button.grid(row=11, column=4, pady=10, padx=10, ipadx=5, ipady=5)

    global bill_button
    bill_button = ttk.Button(window, text="Generate Bill", command=bill)
    bill_button.grid(row=11, column=3, pady=10, padx=10, ipadx=5, ipady=5)


# Second screen
def physical():
    try:
        bill_button.grid_forget()
    except:
        pass

    # delete back button
    try:
        back_to_physical.grid_forget()
        chemical_frame.grid_forget()
        next_half_chemical.grid_forget()
    except:
        pass

    # Checking last entry is Available or not.
    try:
        cur.execute('''SELECT max(id) FROM main_data''')
        max_id = cur.fetchall()[0][0]
    except:
        max_id = None

    global sampleReference
    sampleReference = entry_sampleReference.get() + ', ' + entry_village.get().strip() + ', ' + entry_location.get().strip() + ', ' + entry_sourceType.get()

    try:
        cur.execute(
            '''INSERT INTO main_data(customer, sample_number, sample_drawn_by, sample_drawn_date, sample_reached_lab, test_start_date, test_end_date, sample_reference, village, source_type, location_of_source ) VALUES(?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)''',
            (entry_customer.get(), entry_sampleNumber.get(), entry_sampleDrawn.get(), entry_sampleDrawnDate['text'],
             entry_sampleReached['text'], entry_testStart['text'], entry_testEnd['text'], entry_sampleReference.get(),
             entry_village.get(), entry_sourceType.get(), entry_location.get()))
    except:
        cur.execute('''SELECT id from main_data WHERE sample_number = ?''', (entry_sampleNumber.get(),))
        id_no = (cur.fetchone()[0])
        cur.execute(
            '''UPDATE main_data SET customer=?, sample_number=?, sample_drawn_by=?, sample_drawn_date=?, sample_reached_lab=?, test_start_date=?, test_end_date=?, sample_reference=? WHERE id = ?''',
            (entry_customer.get(), entry_sampleNumber.get(), entry_sampleDrawn.get(), entry_sampleDrawnDate['text'],
             entry_sampleReached['text'], entry_testStart['text'], entry_testEnd['text'], entry_sampleReference.get(),
             id_no))

    connection.commit()

    if max_id is None:
        pass
    else:
        cur.execute('''SELECT * FROM main_data WHERE id = ?''', (max_id,))
        all = cur.fetchall()
        appearance = all[0][12]
        colour = all[0][13]
        turbidity = all[0][15]
        electrical_conductivity = all[0][16]
        TDS = all[0][17]
        total_solids = all[0][18]
        total_suspended_solids = all[0][19]

    cur.execute('''SELECT id from main_data WHERE sample_number = ?''', (entry_sampleNumber.get(),))
    global id
    id = (cur.fetchone()[0])
    cur.execute('''SELECT * FROM main_data''')

    frame.grid_forget()
    next_button.grid_forget()

    global frame2
    frame2 = LabelFrame(window, padx=20, pady=20)
    frame2.grid(padx=20, pady=20, columnspan=5)

    leb14 = Label(frame2, text='1. Physical Parameters')
    leb14.grid(column=0, row=11, pady=10)

    leb15 = Label(frame2, text='Sr. No.', bg='grey')
    leb15.grid(column=0, row=12, sticky=NSEW)

    leb16 = Label(frame2, text='Parameters', bg='grey')
    leb16.grid(column=1, row=12, sticky=NSEW)

    leb17 = Label(frame2, text='Acceptable Limit as per BIS', wraplength=150, bg='grey')
    leb17.grid(column=2, row=12, sticky=NSEW)

    leb18 = Label(frame2, text='Permissible Limit as per BIS', wraplength=150, bg='grey')
    leb18.grid(column=3, row=12, sticky=NSEW)

    leb19 = Label(frame2, text='Results', bg='grey')
    leb19.grid(column=4, row=12, sticky=NSEW)

    # Physical Parameters
    leb20 = Label(frame2, text='1')
    leb20.grid(column=0, row=13)

    # Appearance
    leb21 = Label(frame2, text='Appearance')
    leb21.grid(column=1, row=13, sticky=W)

    leb22 = Label(frame2, text='-')
    leb22.grid(column=2, row=13)

    leb23 = Label(frame2, text='')
    leb23.grid(column=3, row=13)

    global entry_appearance
    entry_appearance = Entry(frame2)
    entry_appearance.grid(column=4, row=13)
    entry_appearance.insert(0, appearance)

    # Colour
    leb24 = Label(frame2, text='2')
    leb24.grid(column=0, row=14)

    leb25 = Label(frame2, text='Colour')
    leb25.grid(column=1, row=14, sticky=W)

    leb26 = Label(frame2, text='-')
    leb26.grid(column=2, row=14)

    leb27 = Label(frame2, text='')
    leb27.grid(column=3, row=14, sticky=W)

    global entry_colour
    entry_colour = Entry(frame2)
    entry_colour.grid(column=4, row=14)
    entry_colour.insert(0, colour)

    # Odour
    leb28 = Label(frame2, text='3')
    leb28.grid(column=0, row=15)

    leb29 = Label(frame2, text='Odour')
    leb29.grid(column=1, row=15, sticky=W)

    leb30 = Label(frame2, text='Unobjectionable')
    leb30.grid(column=2, row=15, sticky=W)

    leb31 = Label(frame2, text='Unobjectionable')
    leb31.grid(column=3, row=15, sticky=W)

    global entry_odour
    entry_odour = ttk.Combobox(frame2, values=['Unobjectionable', 'Objectionable'])
    entry_odour.grid(column=4, row=15)
    entry_odour.current(0)

    # Turbidity NTU
    leb32 = Label(frame2, text='4')
    leb32.grid(column=0, row=16)

    leb33 = Label(frame2, text='Turbidity NTU')
    leb33.grid(column=1, row=16, sticky=W)

    leb34 = Label(frame2, text='1')
    leb34.grid(column=2, row=16)

    leb35 = Label(frame2, text='5')
    leb35.grid(column=3, row=16)

    global entry_turbidity
    entry_turbidity = ttk.Combobox(frame2, values=['Not tested', turbidity])
    entry_turbidity.current(1)
    entry_turbidity.grid(column=4, row=16)

    # Electrical Conductivity (EC)
    leb36 = Label(frame2, text='5')
    leb36.grid(column=0, row=17)

    leb37 = Label(frame2, text='Electrical Conductivity (EC)')
    leb37.grid(column=1, row=17, sticky=W)

    leb38 = Label(frame2, text='-')
    leb38.grid(column=2, row=17)

    leb39 = Label(frame2, text='')
    leb39.grid(column=3, row=17, sticky=W)

    global entry_electricalConductivity
    entry_electricalConductivity = ttk.Combobox(frame2, values=['Not tested', electrical_conductivity])
    entry_electricalConductivity.current(1)
    entry_electricalConductivity.grid(column=4, row=17)

    # Total Dissolved Solids (TDS)
    leb40 = Label(frame2, text='6')
    leb40.grid(column=0, row=18)

    leb41 = Label(frame2, text='Total Dissolved Solids (TDS)')
    leb41.grid(column=1, row=18, sticky=W)

    leb42 = Label(frame2, text='500')
    leb42.grid(column=2, row=18)

    leb43 = Label(frame2, text='2000')
    leb43.grid(column=3, row=18)

    global entry_total_dissolve
    entry_total_dissolve = ttk.Combobox(frame2, values=['Not tested', TDS])
    entry_total_dissolve.current(1)
    entry_total_dissolve.grid(column=4, row=18)

    # ** New addition**
    # Total Solids

    leb44 = Label(frame2, text='7')
    leb44.grid(column=0, row=19)

    leb45 = Label(frame2, text='Total Solids')
    leb45.grid(column=1, row=19, sticky=W)

    leb46 = Label(frame2, text='')
    leb46.grid(column=2, row=19)

    leb47 = Label(frame2, text='')
    leb47.grid(column=3, row=19)

    global entry_total_solids
    entry_total_solids = ttk.Combobox(frame2, values=['Not tested', total_solids])
    entry_total_solids.current(1)
    entry_total_solids.grid(column=4, row=19)

    # ** New addition**
    # Total suspended solids

    leb48 = Label(frame2, text='9')
    leb48.grid(column=0, row=20)

    leb49 = Label(frame2, text='Total Suspended Solids')
    leb49.grid(column=1, row=20, sticky=W)

    leb50 = Label(frame2, text='')
    leb50.grid(column=2, row=20)

    leb51 = Label(frame2, text='')
    leb51.grid(column=3, row=20)

    global entry_total_suspended_solids
    entry_total_suspended_solids = ttk.Combobox(frame2, values=['Not tested', total_suspended_solids])
    entry_total_suspended_solids.current(1)
    entry_total_suspended_solids.grid(column=4, row=20)

    global next_button1
    next_button1 = ttk.Button(window, text="Next", command=chemical)
    next_button1.grid(row=11, column=4, pady=10, padx=10, ipadx=5, ipady=5)

    global back_to_start
    back_to_start = Button(window, text="Back", command=start)
    back_to_start.grid(row=11, column=1, pady=10, padx=10, ipadx=5, ipady=5)


# Third screen
def chemical():
    try:
        back_to_start.grid_forget()
    except:
        pass

    # delete chemical back button
    try:
        back_to_chemical.grid_forget()
        next_button1.grid_forget()
        next_comments.grid_forget()
        chemical_half.grid_forget()
    except:
        pass
    frame2.grid_forget()
    next_button1.grid_forget()

    cur.execute(
        '''UPDATE main_data SET appearance = ?, colour = ?, odour = ?, turbidity = ?, electrical_conductivity = ?, total_dissolved_solids = ?, total_solids = ?, total_suspended_solids = ? WHERE id = ?''',
        (entry_appearance.get(), entry_colour.get(), entry_odour.get(), entry_turbidity.get(),
         entry_electricalConductivity.get(), entry_total_dissolve.get(), entry_total_solids.get(),
         entry_total_suspended_solids.get(), id))

    global chemical_frame
    global next_half_chemical

    cur.execute('''SELECT max(id) FROM main_data''')
    max_id = cur.fetchall()[0][0]
    if max_id - 1 < 1:
        pass
    else:
        cur.execute('''SELECT * FROM main_data WHERE id = ?''', ((max_id - 1),))
        all = cur.fetchall()
        ph = all[0][20]
        ph_alkalinity = all[0][21]
        total_alkalinity = all[0][22]
        total_hardness = all[0][23]
        ca = all[0][24]
        mg = all[0][25]
        fe = all[0][26]
        na = all[0][27]
        k = all[0][28]
        nh3 = all[0][29]
        no2 = all[0][30]
        no3 = all[0][31]

    chemical_frame = LabelFrame(window, padx=20, pady=20)
    chemical_frame.grid(padx=20, pady=20, columnspan=5)

    leb44 = Label(chemical_frame, text='2.1 Chemical Parameters')
    leb44.grid(column=0, row=19, pady=10)

    # Headers
    leb45 = Label(chemical_frame, text='Sr. No.', bg='grey')
    leb45.grid(column=0, row=20, sticky=NSEW)

    leb46 = Label(chemical_frame, text='Parameters', bg='grey')
    leb46.grid(column=1, row=20, sticky=NSEW)

    leb47 = Label(chemical_frame, text='Acceptable Limit as per BIS', wraplength=150, bg='grey')
    leb47.grid(column=2, row=20, sticky=NSEW)

    leb48 = Label(chemical_frame, text='Permissible Limit as per BIS', wraplength=150, bg='grey')
    leb48.grid(column=3, row=20, sticky=NSEW)

    leb49 = Label(chemical_frame, text='Results', bg='grey')
    leb49.grid(column=4, row=20, sticky=NSEW)

    # pH
    leb50 = Label(chemical_frame, text='1')
    leb50.grid(column=0, row=21)

    leb51 = Label(chemical_frame, text='pH')
    leb51.grid(column=1, row=21, sticky=W)

    leb52 = Label(chemical_frame, text='6.5 - 8.5')
    leb52.grid(column=2, row=21)

    leb53 = Label(chemical_frame, text='6.5 - 8.5')
    leb53.grid(column=3, row=21)

    global entry_ph
    entry_ph = ttk.Combobox(chemical_frame, values=['Not tested', ph])
    entry_ph.current(1)
    entry_ph.grid(column=4, row=21)

    # PH Alkalinity as Caco3
    leb54 = Label(chemical_frame, text='2')
    leb54.grid(column=0, row=22)

    leb55 = Label(chemical_frame, text='PH Alkalinity as Caco3')
    leb55.grid(column=1, row=22, sticky=W)

    leb56 = Label(chemical_frame, text='-')
    leb56.grid(column=2, row=22)

    leb57 = Label(chemical_frame, text='-')
    leb57.grid(column=3, row=22)

    global entry_phAlkalinity
    entry_phAlkalinity = ttk.Combobox(chemical_frame, values=['Not tested', ph_alkalinity])
    entry_phAlkalinity.current(1)
    entry_phAlkalinity.grid(column=4, row=22)

    # Total Alkalinity as CaCO3
    leb58 = Label(chemical_frame, text='3')
    leb58.grid(column=0, row=23)

    leb59 = Label(chemical_frame, text='Total Alkalinity as CaCO3')
    leb59.grid(column=1, row=23, sticky=W)

    leb60 = Label(chemical_frame, text='200')
    leb60.grid(column=2, row=23)

    leb61 = Label(chemical_frame, text='600')
    leb61.grid(column=3, row=23)

    global entry_totalAlkalinity
    entry_totalAlkalinity = ttk.Combobox(chemical_frame, values=['Not tested', total_alkalinity])
    entry_totalAlkalinity.current(1)
    entry_totalAlkalinity.grid(column=4, row=23)

    # Total Hardness as CaCO3
    leb62 = Label(chemical_frame, text='4')
    leb62.grid(column=0, row=24)

    leb63 = Label(chemical_frame, text='Total Hardness as CaCO3')
    leb63.grid(column=1, row=24, sticky=W)

    leb64 = Label(chemical_frame, text='200')
    leb64.grid(column=2, row=24)

    leb65 = Label(chemical_frame, text='600')
    leb65.grid(column=3, row=24)

    global entry_totalHardness
    entry_totalHardness = ttk.Combobox(chemical_frame, values=['Not tested', total_hardness])
    entry_totalHardness.current(1)
    entry_totalHardness.grid(column=4, row=24)

    # Calcium as Ca
    leb66 = Label(chemical_frame, text='5')
    leb66.grid(column=0, row=25)

    leb67 = Label(chemical_frame, text='Calcium as Ca')
    leb67.grid(column=1, row=25, sticky=W)

    leb68 = Label(chemical_frame, text='75')
    leb68.grid(column=2, row=25)

    leb69 = Label(chemical_frame, text='100')
    leb69.grid(column=3, row=25)

    global entry_ca
    entry_ca = ttk.Combobox(chemical_frame, values=['Not tested', ca])
    entry_ca.current(1)
    entry_ca.grid(column=4, row=25)

    # Magnesium as Mg
    leb70 = Label(chemical_frame, text='6')
    leb70.grid(column=0, row=26)

    leb71 = Label(chemical_frame, text='Magnesium as Mg')
    leb71.grid(column=1, row=26, sticky=W)

    leb72 = Label(chemical_frame, text='30')
    leb72.grid(column=2, row=26)

    leb73 = Label(chemical_frame, text='150')
    leb73.grid(column=3, row=26)

    global entry_mg
    entry_mg = ttk.Combobox(chemical_frame, values=['Not tested', mg])
    entry_mg.current(1)
    entry_mg.grid(column=4, row=26)

    # Total iron as Fe
    leb74 = Label(chemical_frame, text='7')
    leb74.grid(column=0, row=27)

    leb75 = Label(chemical_frame, text='Total iron as Fe')
    leb75.grid(column=1, row=27, sticky=W)

    leb76 = Label(chemical_frame, text='0.1')
    leb76.grid(column=2, row=27)

    leb77 = Label(chemical_frame, text='1.0')
    leb77.grid(column=3, row=27)

    global entry_fe
    entry_fe = ttk.Combobox(chemical_frame, values=['Not tested', fe])
    entry_fe.current(1)
    entry_fe.grid(column=4, row=27)

    # Sodium as Na
    leb78 = Label(chemical_frame, text='8')
    leb78.grid(column=0, row=28)

    leb79 = Label(chemical_frame, text='Sodium as Na')
    leb79.grid(column=1, row=28, sticky=W)

    leb80 = Label(chemical_frame, text='-')
    leb80.grid(column=2, row=28)

    leb81 = Label(chemical_frame, text='-')
    leb81.grid(column=3, row=28)

    global entry_na
    entry_na = ttk.Combobox(chemical_frame, values=['Not tested', na])
    entry_na.current(1)
    entry_na.grid(column=4, row=28)

    # Potassium as K
    leb82 = Label(chemical_frame, text='9')
    leb82.grid(column=0, row=29)

    leb83 = Label(chemical_frame, text='Potassium as K')
    leb83.grid(column=1, row=29, sticky=W)

    leb84 = Label(chemical_frame, text='-')
    leb84.grid(column=2, row=29)

    leb85 = Label(chemical_frame, text='-')
    leb85.grid(column=3, row=29)

    global entry_k
    entry_k = ttk.Combobox(chemical_frame, values=['Not tested', k])
    entry_k.current(1)
    entry_k.grid(column=4, row=29)

    # Free Ammonia as NH3
    leb86 = Label(chemical_frame, text='10')
    leb86.grid(column=0, row=30)

    leb87 = Label(chemical_frame, text='Free Ammonia as NH3')
    leb87.grid(column=1, row=30, sticky=W)

    leb88 = Label(chemical_frame, text='0.5')
    leb88.grid(column=2, row=30)

    leb89 = Label(chemical_frame, text='0.5')
    leb89.grid(column=3, row=30)

    global entry_nh3
    entry_nh3 = ttk.Combobox(chemical_frame, values=['Not tested', nh3])
    entry_nh3.current(1)
    entry_nh3.grid(column=4, row=30)

    ## **New addition here**

    # Nitrate as NO2
    leb90 = Label(chemical_frame, text='11')
    leb90.grid(column=0, row=31)

    leb91 = Label(chemical_frame, text='Nitrate as NO2')
    leb91.grid(column=1, row=31, sticky=W)

    leb92 = Label(chemical_frame, text='0.5')
    leb92.grid(column=2, row=31)

    leb93 = Label(chemical_frame, text='0.5')
    leb93.grid(column=3, row=31)

    global entry_no2
    entry_no2 = ttk.Combobox(chemical_frame, values=['Not tested', no2])
    entry_no2.current(1)
    entry_no2.grid(column=4, row=31)

    # Nitrate as NO3
    leb94 = Label(chemical_frame, text='12')
    leb94.grid(column=0, row=32)

    leb95 = Label(chemical_frame, text='Nitrate as NO3')
    leb95.grid(column=1, row=32, sticky=W)

    leb96 = Label(chemical_frame, text='45')
    leb96.grid(column=2, row=32)

    leb97 = Label(chemical_frame, text='45')
    leb97.grid(column=3, row=32)

    global entry_no3
    entry_no3 = ttk.Combobox(chemical_frame, values=['Not tested', no3])
    entry_no3.current(1)
    entry_no3.grid(column=4, row=32)

    global back_to_physical
    back_to_physical = Button(window, text="Back", command=physical)
    back_to_physical.grid(row=11, column=1, pady=10, padx=10, ipadx=5, ipady=5)

    next_half_chemical = ttk.Button(window, text="Next", command=half_chemical)
    next_half_chemical.grid(row=11, column=4, pady=10, padx=10, ipadx=5, ipady=5)


# Fourth screen
def half_chemical():
    try:
        back_to_physical.grid_forget()
        next_half_chemical.grid_forget()
    except:
        pass

    cur.execute('''UPDATE "main_data" SET ph = ?, ph_alkalinity = ?, total_alkalinity = ?, total_hardness = ?, ca = ?,
     mg = ?, fe = ?, na = ?, k = ?, nh3 = ?, no2=?, no3=?  WHERE id = ?''',
                (entry_ph.get(), entry_phAlkalinity.get(), entry_totalAlkalinity.get(), entry_totalHardness.get(),
                 entry_ca.get(), entry_mg.get(), entry_fe.get(), entry_na.get(), entry_k.get(), entry_nh3.get(),
                 entry_no2.get(), entry_no3.get(), id))

    chemical_frame.grid_forget()
    next_half_chemical.grid_forget()
    global chemical_half
    global next_comments

    cur.execute('''SELECT max(id) FROM main_data''')
    max_id = cur.fetchall()[0][0]
    cur.execute('''SELECT * FROM main_data WHERE id = ?''', ((max_id - 1),))

    if max_id - 1 < 1:
        pass
    else:
        all = cur.fetchall()
        cl = all[0][32]
        f = all[0][33]
        so4 = all[0][34]
        po4 = all[0][35]
        tids_test = all[0][36]
        do = all[0][37]
        bod = all[0][38]
        cod = all[0][39]
        fecal = all[0][40]
        comments = all[0][41]

    chemical_half = LabelFrame(window, padx=20, pady=20)
    chemical_half.grid(padx=20, pady=20, columnspan=5)

    leb44 = Label(chemical_half, text='2.2 Chemical Parameters')
    leb44.grid(column=0, row=19, pady=10)

    leb45 = Label(chemical_half, text='Sr. No.', bg='grey')
    leb45.grid(column=0, row=20, sticky=NSEW)

    leb46 = Label(chemical_half, text='Parameters', bg='grey')
    leb46.grid(column=1, row=20, sticky=NSEW)

    leb47 = Label(chemical_half, text='Acceptable Limit as per BIS', wraplength=150, bg='grey')
    leb47.grid(column=2, row=20, sticky=NSEW)

    leb48 = Label(chemical_half, text='Permissible Limit as per BIS', wraplength=150, bg='grey')
    leb48.grid(column=3, row=20, sticky=NSEW)

    leb49 = Label(chemical_half, text='Results', bg='grey')
    leb49.grid(column=4, row=20, sticky=NSEW)

    # Chloride as Cl

    leb98 = Label(chemical_half, text='13')
    leb98.grid(column=0, row=33)

    leb99 = Label(chemical_half, text='Chloride as Cl')
    leb99.grid(column=1, row=33, sticky=W)

    leb100 = Label(chemical_half, text='250')
    leb100.grid(column=2, row=33)

    leb101 = Label(chemical_half, text='1000')
    leb101.grid(column=3, row=33)

    global entry_cl
    entry_cl = ttk.Combobox(chemical_half, values=['Not tested', cl])
    entry_cl.current(1)
    entry_cl.grid(column=4, row=33)

    # Fluoride as F

    leb102 = Label(chemical_half, text='14')
    leb102.grid(column=0, row=34)

    leb103 = Label(chemical_half, text='Fluoride as F')
    leb103.grid(column=1, row=34, sticky=W)

    leb104 = Label(chemical_half, text='1.0')
    leb104.grid(column=2, row=34)

    leb105 = Label(chemical_half, text='1.5')
    leb105.grid(column=3, row=34)

    global entry_f
    entry_f = ttk.Combobox(chemical_half, values=['Not tested', f])
    entry_f.current(1)
    entry_f.grid(column=4, row=34)

    # Sulphate as SO4

    leb106 = Label(chemical_half, text='15')
    leb106.grid(column=0, row=35)

    leb107 = Label(chemical_half, text='Sulphate as SO4')
    leb107.grid(column=1, row=35, sticky=W)

    leb108 = Label(chemical_half, text='200')
    leb108.grid(column=2, row=35)

    leb109 = Label(chemical_half, text='400')
    leb109.grid(column=3, row=35)

    global entry_so4
    entry_so4 = ttk.Combobox(chemical_half, values=['Not tested', so4])
    entry_so4.current(1)
    entry_so4.grid(column=4, row=35)

    # Phosphate as PO4

    leb110 = Label(chemical_half, text='16')
    leb110.grid(column=0, row=36)

    leb111 = Label(chemical_half, text='Phosphate as PO4')
    leb111.grid(column=1, row=36, sticky=W)

    leb112 = Label(chemical_half, text='0.5')
    leb112.grid(column=2, row=36)

    leb113 = Label(chemical_half, text='0.5')
    leb113.grid(column=3, row=36)

    global entry_po4
    entry_po4 = ttk.Combobox(chemical_half, values=['Not tested', po4])
    entry_po4.current(1)
    entry_po4.grid(column=4, row=36)

    # Tids Test 4 hours as O

    leb114 = Label(chemical_half, text='17')
    leb114.grid(column=0, row=37)

    leb115 = Label(chemical_half, text='Tids Test 4 hours as O')
    leb115.grid(column=1, row=37, sticky=W)

    leb116 = Label(chemical_half, text='-')
    leb116.grid(column=2, row=37)

    leb117 = Label(chemical_half, text='-')
    leb117.grid(column=3, row=37)

    global entry_O
    entry_O = ttk.Combobox(chemical_half, values=['Not tested', tids_test])
    entry_O.current(1)
    entry_O.grid(column=4, row=37)

    # Dissolved oxygen
    leb118 = Label(chemical_half, text='18')
    leb118.grid(column=0, row=38)

    leb119 = Label(chemical_half, text='Dissolved oxygen (DO)')
    leb119.grid(column=1, row=38, sticky=W)

    leb120 = Label(chemical_half, text='-')
    leb120.grid(column=2, row=38)

    leb120 = Label(chemical_half, text='-')
    leb120.grid(column=3, row=38)

    global entry_do
    entry_do = ttk.Combobox(chemical_half, values=['Not tested', do])
    entry_do.current(1)
    entry_do.grid(column=4, row=38)

    # Biological Oxygen Demand (BOD)

    leb121 = Label(chemical_half, text='19')
    leb121.grid(column=0, row=39)

    leb122 = Label(chemical_half, text='Biological Oxygen Demand (BOD)')
    leb122.grid(column=1, row=39, sticky=W)

    leb123 = Label(chemical_half, text='-')
    leb123.grid(column=2, row=39)

    leb124 = Label(chemical_half, text='-')
    leb124.grid(column=3, row=39)

    global entry_bod
    entry_bod = ttk.Combobox(chemical_half, values=['Not tested', bod])
    entry_bod.current(1)
    entry_bod.grid(column=4, row=39)

    # **New addition
    # 1) Dissolved oxygen row number 38 above BOD
    # 2) Chemical oxygen demand
    leb125 = Label(chemical_half, text='20')
    leb125.grid(column=0, row=40)

    leb126 = Label(chemical_half, text='Chemical Oxygen Demand (COD)')
    leb126.grid(column=1, row=40, sticky=W)

    leb127 = Label(chemical_half, text='-')
    leb127.grid(column=2, row=40)

    leb128 = Label(chemical_half, text='-')
    leb128.grid(column=3, row=40)

    global entry_cod
    entry_cod = ttk.Combobox(chemical_half, values=['Not tested', cod])
    entry_cod.current(1)
    entry_cod.grid(column=4, row=40)

    leb129 = Label(chemical_half, text='3. Bacteriological Parameters')
    leb129.grid(column=0, row=41, pady=10)

    # Headers
    # leb130 = Label(chemical_half, text='Sr. No.', bg='grey')
    # leb130.grid(column=0, row=42, sticky=NSEW)

    # leb131 = Label(chemical_half, text='Parameters', bg='grey')
    # leb131.grid(column=1, row=42, sticky=NSEW)

    # leb132 = Label(chemical_half, text='Acceptable Limit as per BIS', wraplength=150, bg='grey')
    # leb132.grid(column=2, row=42, sticky=NSEW)

    # leb133 = Label(chemical_half, text='Permissible Limit as per BIS', wraplength=150, bg='grey')
    # leb133.grid(column=3, row=42, sticky=NSEW)

    # leb134 = Label(chemical_half, text='Results', bg='grey')
    # leb134.grid(column=4, row=42, sticky=NSEW)

    # Fecal coliform entry
    leb135 = Label(chemical_half, text='1')
    leb135.grid(column=0, row=43)

    leb136 = Label(chemical_half, text='Faecal Coliform')
    leb136.grid(column=1, row=43, sticky=W)

    leb137 = Label(chemical_half, text='-')
    leb137.grid(column=2, row=43)

    leb138 = Label(chemical_half, text='-')
    leb138.grid(column=3, row=43)

    global entry_fecal_coliform
    entry_fecal_coliform = ttk.Combobox(chemical_half, values=['Not tested', fecal])
    entry_fecal_coliform.current(1)
    entry_fecal_coliform.grid(column=4, row=43)

    comment = Label(chemical_half, text='3. Comments/ Remarks', width=40)
    comment.grid(column=0, row=44, pady=10, sticky=E)

    global comment_entry
    comment_entry = Entry(chemical_half, text="comment")
    comment_entry.grid(column=0, row=44, columnspan=5, ipady=30, padx=50)
    comment_entry.delete(0, END)
    comment_entry.insert(0, comments)

    leb139 = Label(chemical_half, text="Date:")
    leb139.grid(column=0, row=45, sticky=E, pady=10)

    global button_date
    button_date = ttk.Button(chemical_half, text="Select a date", command=lambda: datePicker('main_date'))
    button_date.grid(column=1, row=45, sticky=W, pady=10)
    button_date['text'] = date

    global back_to_chemical
    back_to_chemical = Button(window, text="Back", command=chemical)
    back_to_chemical.grid(row=11, column=1, pady=10, padx=10, ipadx=5, ipady=5)

    next_comments = ttk.Button(window, text="Save", command=confirm)
    next_comments.grid(row=11, column=4, pady=10, padx=10, ipadx=5, ipady=5)


# Confirmation pop up
def confirm():
    cur.execute(
        '''UPDATE main_data SET cl=?, f=?, so4=?, po4=?, tids_test=?, do=?, bod=?, chemical_oxygen_demand=?, fecal_coliform=?, comments=?, date=?  WHERE id = ?''',
        (entry_cl.get(), entry_f.get(), entry_so4.get(), entry_po4.get(),
         entry_O.get(), entry_do.get(), entry_bod.get(), entry_cod.get(), entry_fecal_coliform.get(),
         comment_entry.get(), button_date['text'], id))
    cur.execute('''SELECT max(id) FROM main_data''')
    max_id = cur.fetchall()[0][0]
    cur.execute('''SELECT * FROM main_data WHERE id = ?''', (max_id,))
    global top
    top = Toplevel(window, )

    label_list = ['id', 'Customer Details: ', 'Sample Number: ', 'Sample Drawn by: ', 'Sample Drawn Date ',
                  'Sample Reached Lab on:', 'Test Start Date', 'Test End Date', 'Sample Reference', 'Village',
                  'Source type', 'Location of source:', 'Appearance:', 'Colour:', 'Odour:', 'Turbidity NTU:',
                  'Electrical Conductivity (EC):', 'Total Dissolved Solids (TDS):', 'Total Solids:',
                  'Total Suspended Solids:', 'pH:', 'Ph Alkalinity as Caco3:',
                  'Total Alkalinity as CaCO3:', 'Total Hardness as CaCO3:', 'Calcium as Ca:', 'Magnesium as Mg:',
                  'Total iron as Fe:', 'Sodium as Na:', 'Potassium as K:', 'Free Ammonia as NH3:', 'Nitrite as NO2:',
                  'Nitrate as NO3:', 'Chloride as Cl:', 'Fluoride as F:', 'Sulphate as SO4:', 'Phosphate as PO4:',
                  'Tids test 4 Hours as O:',
                  'Dissolved Oxygen (DO):', 'Biological Oxygen Demand (BOD):', 'Chemical Oxygen Demand:',
                  'Fecal Coliform:', 'Comments / Remarks', 'Date:']
    verification_frame = LabelFrame(top, text="Confirmation", padx=20, pady=20, height=50)
    verification_frame.grid(padx=10, pady=10, columnspan=2)
    # scrollbar = Scrollbar(top)
    # scrollbar.grid(column =3, rowspan=18 )

    cur.execute('''SELECT * FROM main_data WHERE id = ?''', (max_id,))
    verification_data = cur.fetchall()[0]

    x, y, l, count = 0, 0, 0, 1
    for i in verification_data:
        if x <= 21:
            very_label = Label(verification_frame, text=label_list[l] + ' ' + str(i))
            very_label.grid(row=x, column=0, sticky=W)
            x = x + 1
        if x > 21:
            very_label = Label(verification_frame, text=label_list[l] + ' ' + str(i))
            very_label.grid(row=y, column=1, sticky=W)
            y = y + 1
        count = count + 1
        l = l + 1

    not_confirmation_button = ttk.Button(top, text="Back", command=top_exit)
    not_confirmation_button.grid(row=1, column=0, pady=10)
    confirmation_button = ttk.Button(top, text='Confirm', command=saveData)
    confirmation_button.grid(row=1, column=1, pady=10)


# Final data saving function - sqlite connection commit()
def saveData():
    connection.commit()
    saveWord()
    messagebox.showinfo("Success", "Your Data has been saved successfully")
    back_to_chemical.grid_forget()
    next_comments.grid_forget()
    start()


# Command only for exit the popup window
def top_exit():
    top.destroy()


def saveWord():
    cur.execute('''SELECT * FROM main_data WHERE id=?''', (id,))
    all = cur.fetchall()[0]
    doc = Document('sample.docx')
    doc.tables[0].cell(0, 1).text = all[1]
    doc.tables[0].rows[0].cells[1].paragraphs[0].runs[0].font.name = 'Footlight MT'
    doc.tables[0].cell(1, 1).text = str(all[2])
    doc.tables[0].rows[1].cells[1].paragraphs[0].runs[0].font.name = 'Footlight MT'
    doc.tables[0].cell(1, 3).text = all[3]
    doc.tables[0].rows[1].cells[3].paragraphs[0].runs[0].font.name = 'Footlight MT'
    doc.tables[0].cell(2, 1).text = all[4]
    doc.tables[0].rows[2].cells[1].paragraphs[0].runs[0].font.name = 'Footlight MT'
    doc.tables[0].cell(2, 3).text = all[5]
    doc.tables[0].rows[2].cells[3].paragraphs[0].runs[0].font.name = 'Footlight MT'
    doc.tables[0].cell(3, 1).text = all[6]
    doc.tables[0].rows[3].cells[1].paragraphs[0].runs[0].font.name = 'Footlight MT'
    doc.tables[0].cell(3, 3).text = all[7]
    doc.tables[0].rows[3].cells[3].paragraphs[0].runs[0].font.name = 'Footlight MT'
    doc.tables[0].cell(4, 1).text = str(sampleReference)
    doc.tables[0].rows[4].cells[1].paragraphs[0].runs[0].font.name = 'Footlight MT'

    global exceed
    exceed =''

    x, y = 1, 12
    for cell in range(8):
        doc.tables[1].cell(x, 4).text = str(all[y])
        doc.tables[1].rows[x].cells[4].paragraphs[0].runs[0].font.name = 'Footlight MT'
        x += 1
        y += 1
    # Table 1 Limit border bold
    try:
        if doc.tables[1].cell(3, 4).text == 'Objectionable':
            exceed = doc.tables[1].cell(3, 1).text
            doc.tables[1].rows[3].cells[4].paragraphs[0].runs[0].font.bold = True
    except:
        pass

    try:
        if float(doc.tables[1].cell(4, 4).text) > 5:
            exceed = exceed + ', ' + doc.tables[1].cell(4, 1).text
            doc.tables[1].rows[4].cells[4].paragraphs[0].runs[0].font.bold = True
    except:
        pass

    try:
        if float(doc.tables[1].cell(6, 4).text) > 2000:
            exceed = exceed + ', ' + doc.tables[1].cell(6, 1).text
            doc.tables[1].rows[6].cells[4].paragraphs[0].runs[0].font.bold = True
    except:
        pass

    x, y = 1, 20
    for cell in range(20):
        doc.tables[2].cell(x, 4).text = str(all[y])
        doc.tables[2].rows[x].cells[4].paragraphs[0].runs[0].font.name = 'Footlight MT'
        x += 1
        y += 1

    # Table 2 limit bold
    # ph
    try:
        if 6.5 > float(doc.tables[2].cell(1, 4).text) or float(doc.tables[2].cell(4, 4).text) > 8.5:
            exceed = exceed + ', ' + doc.tables[2].cell(1, 1).text
            doc.tables[2].rows[1].cells[4].paragraphs[0].runs[0].font.bold = True
    except:
        pass

    # Total Alkalinity
    try:
        if float(doc.tables[2].cell(3, 4).text) > 600:
            exceed = exceed + ', ' + doc.tables[2].cell(3, 1).text
             
            doc.tables[2].rows[3].cells[4].paragraphs[0].runs[0].font.bold = True
    except:
        pass

    # Total Hardness
    try:
        if float(doc.tables[2].cell(4, 4).text) > 600:
            exceed = exceed + ', ' + doc.tables[2].cell(4, 1).text
             
            doc.tables[2].rows[4].cells[4].paragraphs[0].runs[0].font.bold = True
    except:
        pass
    # Ca
    try:
        if float(doc.tables[2].cell(5, 4).text) > 100:
            exceed = exceed + ', ' + doc.tables[2].cell(5, 1).text
             
            doc.tables[2].rows[5].cells[4].paragraphs[0].runs[0].font.bold = True
    except:
        pass

    # Mg
    try:
        if float(doc.tables[2].cell(6, 4).text) > 150:
            exceed = exceed + ', ' + doc.tables[2].cell(6, 1).text
             
            doc.tables[2].rows[6].cells[4].paragraphs[0].runs[0].font.bold = True
    except:
        pass

    # Fe
    try:
        if float(doc.tables[2].cell(7, 4).text) > 1.0:
            exceed = exceed + ', ' + doc.tables[2].cell(7, 1).text
            doc.tables[2].rows[7].cells[4].paragraphs[0].runs[0].font.bold = True
    except:
        pass
    # NH3
    try:
        if float(doc.tables[2].cell(10, 4).text) > 0.5:
            exceed = exceed + ', ' + doc.tables[2].cell(10, 1).text
             
            doc.tables[2].rows[10].cells[4].paragraphs[0].runs[0].font.bold = True
    except:
        pass
    # NO2
    try:
        if float(doc.tables[2].cell(11, 4).text) > 0.5:
            exceed = exceed + ', ' + doc.tables[2].cell(11, 1).text
             
            doc.tables[2].rows[11].cells[4].paragraphs[0].runs[0].font.bold = True
    except:
        pass
    # NO3
    try:
        if float(doc.tables[2].cell(12, 4).text) > 45:
            exceed = exceed + ', ' + doc.tables[2].cell(12, 1).text
             
            doc.tables[2].rows[12].cells[4].paragraphs[0].runs[0].font.bold = True
    except:
        pass
    # Cl
    try:
        if float(doc.tables[2].cell(13, 4).text) > 1000:
            exceed = exceed + ', ' + doc.tables[2].cell(13, 1).text
             
            doc.tables[2].rows[13].cells[4].paragraphs[0].runs[0].font.bold = True
    except:
        pass
    # Fluoride as F
    try:
        if float(doc.tables[2].cell(14, 4).text) > 1.5:
            exceed = exceed + ', ' + doc.tables[2].cell(14, 1).text
             
            doc.tables[2].rows[14].cells[4].paragraphs[0].runs[0].font.bold = True
    except:
        pass
    # Sulphate as SO4
    try:
        if float(doc.tables[2].cell(15, 4).text) > 400:
            exceed = exceed + ', ' + doc.tables[2].cell(15, 1).text
             
            doc.tables[2].rows[15].cells[4].paragraphs[0].runs[0].font.bold = True
    except:
        pass

    # Phosphate as PO4
    try:
        if float(doc.tables[2].cell(16, 4).text) > 0.5:
            exceed = exceed + ', ' + doc.tables[2].cell(16, 1).text
             
            doc.tables[2].rows[16].cells[4].paragraphs[0].runs[0].font.bold = True
    except:
        pass

    doc.tables[3].cell(1, 4).text = str(all[40])
    doc.tables[3].rows[1].cells[4].paragraphs[0].runs[0].font.name = 'Footlight MT'

    # if doc.tables[3].cell(1, 4).text == 'Not tested':
    #   doc.tables[3].rows[1].cells[4].paragraphs[0].runs[0].font.bold = True
    if exceed:
        doc.tables[4].cell(0, 0).text = exceed + '' + " are exiding limit " + str(all[41])
    else:
        doc.tables[4].cell(0, 0).text = str(all[41])

    doc.tables[4].rows[0].cells[0].paragraphs[0].runs[0].font.name = 'Footlight MT'

    filename = str(all[1]).strip() + '_sample_no_' + str(all[2]).strip() + '.docx'

    for paragraph in doc.paragraphs:
        if 'Date' in paragraph.text:
            paragraph.text  =''
            run = paragraph.add_run()
            run.text = "Date: "
            run.bold = True
            run.font.name = 'Footlight MT'
            run1 = paragraph.add_run()
            run1.text = button_date['text']
            run1.font.name = 'Footlight MT'

    doc.save("word/%s" % filename)


date_str = str(datetime.date.today())


def generate():
    from_sample = from_entry.get()
    to_sample = to_entry.get()
    cur.execute('''SELECT count(sample_number) FROM main_data WHERE sample_number BETWEEN ? and ? ''',
                (from_sample, to_sample,))
    x = cur.fetchall()
    total_samples = x[0][0]
    cur.execute('''SELECT customer, sample_drawn_by FROM main_data WHERE sample_number = ?  ''', (to_sample,))
    raw_data = cur.fetchall()
    customer_name = raw_data[0][0]
    sample_drawn_by = raw_data[0][1]

    cur.execute("SELECT * FROM main_data WHERE sample_number = ?", (to_sample,))
    all = cur.fetchall()

    physical_parameter = [all[0][12], all[0][13], all[0][14], all[0][15], all[0][16], all[0][17], all[0][18],
                          all[0][19]]
    physical_tested_parameter_count = 8
    for parameters in physical_parameter:
        if parameters == "Not tested":
            physical_tested_parameter_count -= 1

    chemical_tested_parameter_count = 20
    chemical_parameter = [all[0][20], all[0][21], all[0][22], all[0][23], all[0][24], all[0][25], all[0][26],
                          all[0][27], all[0][28], all[0][29], all[0][30], all[0][31],
                          all[0][32], all[0][33], all[0][34], all[0][35], all[0][36], all[0][37], all[0][38],
                          all[0][39]]
    for parameters in chemical_parameter:
        if parameters == "Not tested":
            chemical_tested_parameter_count -= 1

    if all[0][40] == "Not tested":
        biological_tested_parameter_count = 0
    else:
        biological_tested_parameter_count = 1

    if deduction.get() == '' or None:
        deductionRs = 0
    else:
        deductionRs = int(deduction.get())

    if price_per_sample.get() == 'Other':
        total_amount = (int(total_samples) * int(other_entry.get())) - (int(total_samples) * int(deductionRs))
    else:
        total_amount = (int(total_samples) * int(price_per_sample.get())) - (int(total_samples) * int(deductionRs))

    if TA_sample_collection.get() == '' or TA_sample_collection.get() is None:
        sample_collection = 0
    else:
        sample_collection = int(TA_sample_collection.get())
    grand_total = total_amount + sample_collection
    try:
        doc = Document('water lab bill.docx')
        doc.tables[0].cell(0, 1).text = customer_name
        doc.tables[0].rows[0].cells[1].paragraphs[0].runs[0].font.name = 'Footlight MT'
        if from_sample == to_sample:
            doc.tables[0].cell(1, 1).text = from_sample
        else:
            doc.tables[0].cell(1, 1).text = from_sample + ' ' + 'to' + ' ' + to_sample

        doc.tables[0].rows[1].cells[1].paragraphs[0].runs[0].font.name = 'Footlight MT'

        doc.tables[0].cell(2, 1).text = sample_drawn_by
        doc.tables[0].rows[2].cells[1].paragraphs[0].runs[0].font.name = 'Footlight MT'
        doc.tables[0].cell(3, 1).text = billing_date
        doc.tables[0].rows[3].cells[1].paragraphs[0].runs[0].font.name = 'Footlight MT'
        if billing_rate_category.get() == 'Other':
            doc.tables[0].cell(5, 1).text = other_billing_rate_category.get()
        else:
            doc.tables[0].cell(5, 1).text = billing_rate_category.get()
        doc.tables[0].rows[5].cells[1].paragraphs[0].runs[0].font.name = 'Footlight MT'

        doc.tables[0].cell(4, 2).text = str(physical_tested_parameter_count)
        doc.tables[0].rows[4].cells[2].paragraphs[0].runs[0].font.name = 'Footlight MT'
        doc.tables[0].cell(4, 4).text = str(chemical_tested_parameter_count)
        doc.tables[0].rows[4].cells[4].paragraphs[0].runs[0].font.name = 'Footlight MT'
        doc.tables[0].cell(4, 6).text = str(biological_tested_parameter_count)
        doc.tables[0].rows[4].cells[6].paragraphs[0].runs[0].font.name = 'Footlight MT'

        if billing_rate_category.get() == 'Other':
            doc.tables[0].cell(6, 1).text = str(other_entry.get())
        else:
            doc.tables[0].cell(6, 1).text = str(price_per_sample.get())
        doc.tables[0].rows[6].cells[1].paragraphs[0].runs[0].font.name = 'Footlight MT'
        doc.tables[0].cell(7, 1).text = str(deductionRs)
        doc.tables[0].rows[7].cells[1].paragraphs[0].runs[0].font.name = 'Footlight MT'
        doc.tables[0].cell(8, 1).text = str(total_samples)
        doc.tables[0].rows[8].cells[1].paragraphs[0].runs[0].font.name = 'Footlight MT'
        doc.tables[0].cell(9, 1).text = "Rs.%d" % int(total_amount)
        doc.tables[0].rows[9].cells[1].paragraphs[0].runs[0].font.name = 'Footlight MT'
        if sample_collection == 0:
            doc.tables[0].cell(10, 1).text = '0'
        else:
            doc.tables[0].cell(10, 1).text = ("Rs." + str(TA_sample_collection.get()))
        doc.tables[0].rows[10].cells[1].paragraphs[0].runs[0].font.name = 'Footlight MT'
        doc.tables[0].cell(11, 1).text = "Rs.%d" % grand_total
        doc.tables[0].rows[11].cells[1].paragraphs[0].runs[0].font.name = 'Footlight MT'
        run = doc.paragraphs[27].add_run()
        run.text = 'Date: '
        run.bold = True
        run.font.name = 'Footlight MT'
        run1 = doc.paragraphs[27].add_run()
        run1.text = billing_date
        run1.font.name = 'Footlight MT'
        #doc.paragraphs[27].add_run = "Date:" + ' ' + billing_date

        doc.save(f"word/bill/{from_sample} to {to_sample} bill.docx")
        messagebox.showinfo("Success", "Bill has been generated successfully")
    except:
        messagebox.showerror("Error", "Data has been not saved properly Please enter again or Contact administrator")


def bill():
    billing_window = Toplevel(window)
    billing_window.title('Bills')
    billing_frame = LabelFrame(billing_window, text='Bill', padx=5, pady=5)
    billing_frame.grid(padx=10, pady=10)
    title = Label(billing_frame, text="Billing Window")
    title.grid(row=0, column=0, columnspan=4)
    title.configure(font="Verdana 20")

    select_samples = Label(billing_frame, text="Select the sample numbers of which you want to generate Bill.",
                           fg="red")
    select_samples.grid(row=1, column=0, columnspan=4, ipady=10, ipadx=10)

    from_label = Label(billing_frame, text="From")
    from_label.grid(row=2, column=0, sticky=E)

    global from_entry
    from_entry = Entry(billing_frame, )
    from_entry.grid(row=2, column=1, sticky=W)

    to_label = Label(billing_frame, text="To")
    to_label.grid(row=2, column=2, sticky=E)

    global to_entry
    to_entry = Entry(billing_frame)
    to_entry.grid(row=2, column=3, sticky=W)

    cur.execute('''SELECT max(sample_number) FROM main_data''')
    to_sample = cur.fetchall()[0][0]

    to_entry.insert(0, to_sample)

    billing_rate_category_label = Label(billing_frame, text='Select billing rate category:')
    billing_rate_category_label.grid(row=3, column=0, sticky=E)

    global billing_rate_category
    billing_rate_category = ttk.Combobox(billing_frame, width=70,
                                         values=['Corporate - Category A', 'Projects for Public Cause - Category B',
                                                 'For Government Organisations - Category C',
                                                 'For the individual, DHAN members, NGOs, Farmers, Community and Students - Category D',
                                                 'Other'])
    billing_rate_category.grid(row=3, column=1, columnspan=4, pady=5)
    billing_rate_category.current(3)

    other_price = Label(billing_frame, text="If other:")
    other_price.grid(row=4, column=0, sticky=E)

    global other_billing_rate_category

    other_billing_rate_category = Entry(billing_frame)
    other_billing_rate_category.grid(row=4, column=1, sticky=W)

    price_label = Label(billing_frame, text='Price per sample:')
    price_label.grid(row=5, column=0, sticky=E)

    global price_per_sample
    price_per_sample = ttk.Combobox(billing_frame, value=['1500', '1000', '800', '600', 'Other'])
    price_per_sample.grid(row=5, column=1, sticky=W, pady=5)
    price_per_sample.current(3)

    other_label = Label(billing_frame, text='If Other')
    other_label.grid(row=5, column=2, sticky=E)

    global other_entry
    other_entry = Entry(billing_frame)
    other_entry.grid(row=5, column=3, sticky=W)

    deduction_label = Label(billing_frame, text='Deduction per sample:')
    deduction_label.grid(row=6, column=0, sticky=E)

    global deduction
    deduction = Entry(billing_frame)
    deduction.grid(row=6, column=1, sticky=W)

    TA_sample_collection_label = Label(billing_frame, text="Travel Allowance and sample collection")
    TA_sample_collection_label.grid(row=7, column=0, sticky=E)

    # noinspection PyGlobalUndefined
    global TA_sample_collection
    TA_sample_collection = Entry(billing_frame)
    TA_sample_collection.grid(row=7, column=1, sticky=W)

    global bill_date
    bill_date = ttk.Button(billing_frame, text='-- Select a date --', command=lambda: datePicker('bill_date'))
    bill_date.grid(row=8, column=1, sticky=W)

    generator_button = ttk.Button(billing_frame, text="Generate Bill", command=generate)
    generator_button.grid(row=8, column=3)


def datePicker(button_name):
    def printDate():
        selectedDate = cal.selection_get()
        if button_name == 'drawnDate':
            entry_sampleDrawnDate.configure(text=selectedDate.strftime('%d %B %Y'))
        elif button_name == 'reachedDate':
            entry_sampleReached.configure(text=selectedDate.strftime('%d %B %Y'))
        elif button_name == 'testStart':
            entry_testStart.configure(text=selectedDate.strftime('%d %B %Y'))
        elif button_name == 'testEnd':
            entry_testEnd.configure(text=selectedDate.strftime('%d %B %Y'))
        elif button_name == 'main_date':
            button_date.configure(text=selectedDate.strftime('%d %B %Y'))
        elif button_name == 'bill_date':
            global billing_date
            billing_date = selectedDate.strftime('%d %B %Y')
            bill_date.configure(text=selectedDate.strftime('%d %B %Y'))

        dateWindow.destroy()

    dateWindow = Toplevel(window)
    year, month, day = date.year, date.month, date.day
    cal = Calendar(dateWindow, font="Arial 14", selectmode='day', year=year, month=month, day=day)
    cal.pack(fill="both", expand=True)
    ttk.Button(dateWindow, text="Select", command=printDate).pack(pady=5)


s = ttk.Style(window)
s.theme_use('clam')

start()
window.mainloop()
